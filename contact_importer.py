"""
contact_importer.py
--------------------
Imports YOUR OWN existing contact/client lists from PDF, Word (.docx),
plain text, or CSV files, and merges them into leads.json using the same
schema OutreachAgent/LeadScoutAgent already expect.

This is for lists you already own (existing clients, a contact list you
maintain yourself) — not for scraping new individuals' personal data.

Supported formats:
  - .csv   : columns are mapped by header name (business_name, email, phone, etc.)
  - .txt   : one contact per line, fields separated by comma or tab, OR
             free text where we regex out emails/phones per line
  - .docx  : reads text + any tables; tables are parsed like CSV, plain
             paragraphs are scanned with the same regex fallback as .txt
  - .pdf   : extracts text (and tables where possible) and applies the
             same parsing as .docx/.txt

Install once on your machine:
    pip install python-docx pypdf pdfplumber
"""

import os
import re
import csv
import json
import logging
from io import StringIO

logger = logging.getLogger(__name__)

EMAIL_RE = re.compile(r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}")
PHONE_RE = re.compile(r"(\+?\d{1,3}[\s.\-]?)?(\(?\d{2,4}\)?[\s.\-]?){1,4}\d{3,4}")

LEADS_FILE = "leads.json"

# Recognized header aliases -> canonical lead field name
HEADER_ALIASES = {
    "name": "business_name", "company": "business_name", "business": "business_name",
    "business_name": "business_name", "client": "business_name", "client_name": "business_name",
    "email": "business_email", "business_email": "business_email", "e-mail": "business_email",
    "phone": "phone", "phone_number": "phone", "contact_number": "phone", "tel": "phone",
    "website": "website", "site": "website", "url": "website",
    "address": "address", "location": "address",
    "niche": "niche", "industry": "niche", "category": "niche",
    "notes": "notes",
}


def normalize_row(row: dict) -> dict:
    """Map arbitrary header names to canonical lead fields."""
    lead = {}
    for key, value in row.items():
        if value is None:
            continue
        canonical = HEADER_ALIASES.get(key.strip().lower().replace(" ", "_"))
        if canonical:
            lead[canonical] = str(value).strip()
    return lead


def extract_fallback(line: str) -> dict:
    """For free-text lines with no clear column structure, pull out an
    email and phone if present, and use the remaining text as the name."""
    lead = {}
    email_match = EMAIL_RE.search(line)
    if email_match:
        lead["business_email"] = email_match.group(0)

    # Search for phone candidates only in the text segments NOT already
    # matched as the email, otherwise the phone regex can accidentally
    # latch onto digits inside an email-adjacent segment.
    remainder_for_phone = line
    if email_match:
        remainder_for_phone = remainder_for_phone.replace(email_match.group(0), " ")

    phone_match = PHONE_RE.search(remainder_for_phone)
    if phone_match and len(re.sub(r"\D", "", phone_match.group(0))) >= 7:
        lead["phone"] = phone_match.group(0).strip()

    # Whatever text remains after stripping the matched email/phone is a
    # reasonable guess at the name.
    remainder = line
    if email_match:
        remainder = remainder.replace(email_match.group(0), " ")
    if phone_match and lead.get("phone"):
        remainder = remainder.replace(phone_match.group(0), " ")
    # Collapse stray separators (-, |, ;, multiple spaces) left behind
    # once the email/phone tokens are removed.
    remainder = re.sub(r"[,;|\t]+", " ", remainder)
    remainder = re.sub(r"\s*-\s*", " ", remainder)
    remainder = re.sub(r"\s{2,}", " ", remainder).strip(" -:\t")
    if remainder:
        lead["business_name"] = remainder

    return lead


def parse_csv_text(text: str) -> list:
    leads = []
    reader = csv.DictReader(StringIO(text))
    if reader.fieldnames is None:
        return leads
    for row in reader:
        lead = normalize_row(row)
        if lead.get("business_email") or lead.get("phone"):
            leads.append(lead)
    return leads


def parse_plain_lines(text: str) -> list:
    """Used for .txt content and as a fallback for docx/pdf paragraphs."""
    leads = []
    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue

        # Try comma/tab-delimited first (e.g. "Acme Inc, info@acme.com, 555-1234")
        if "," in line or "\t" in line:
            delimiter = "\t" if "\t" in line else ","
            parts = [p.strip() for p in line.split(delimiter)]
            lead = {}
            for part in parts:
                if EMAIL_RE.fullmatch(part):
                    lead["business_email"] = part
                elif PHONE_RE.fullmatch(part) and len(re.sub(r"\D", "", part)) >= 7:
                    lead["phone"] = part
                elif part and "business_name" not in lead:
                    lead["business_name"] = part
            if lead.get("business_email") or lead.get("phone"):
                leads.append(lead)
                continue

        # Fallback: regex scan the whole line
        lead = extract_fallback(line)
        if lead.get("business_email") or lead.get("phone"):
            leads.append(lead)

    return leads


def parse_docx_tables(path: str) -> list:
    try:
        import docx
    except ImportError:
        raise ImportError("python-docx is required for .docx import. Run: pip install python-docx")

    doc = docx.Document(path)
    leads = []

    for table in doc.tables:
        if not table.rows:
            continue
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        for row in table.rows[1:]:
            values = [cell.text.strip() for cell in row.cells]
            row_dict = dict(zip(headers, values))
            lead = normalize_row(row_dict)
            if lead.get("business_email") or lead.get("phone"):
                leads.append(lead)

    # Also scan plain paragraphs (covers lists not in table form)
    paragraph_text = "\n".join(p.text for p in doc.paragraphs)
    leads.extend(parse_plain_lines(paragraph_text))

    return leads


def parse_pdf(path: str) -> list:
    try:
        import pdfplumber
    except ImportError:
        pdfplumber = None

    leads = []

    if pdfplumber is not None:
        try:
            with pdfplumber.open(path) as pdf:
                for page in pdf.pages:
                    for table in (page.extract_tables() or []):
                        if not table or len(table) < 2:
                            continue
                        headers = [str(h or "").strip() for h in table[0]]
                        for row in table[1:]:
                            row_dict = dict(zip(headers, [str(c or "").strip() for c in row]))
                            lead = normalize_row(row_dict)
                            if lead.get("business_email") or lead.get("phone"):
                                leads.append(lead)
                    text = page.extract_text() or ""
                    leads.extend(parse_plain_lines(text))
            if leads:
                return leads
        except Exception as e:
            logger.warning("pdfplumber failed (%s), falling back to pypdf text-only extraction.", e)

    try:
        from pypdf import PdfReader
    except ImportError:
        raise ImportError("pypdf is required for .pdf import. Run: pip install pypdf")

    reader = PdfReader(path)
    full_text = "\n".join((page.extract_text() or "") for page in reader.pages)
    leads.extend(parse_plain_lines(full_text))
    return leads


def import_contacts(path: str, default_niche: str = "", default_preferred_channel: str = "email") -> dict:
    """
    Parse the file at `path` and merge any found contacts into leads.json.

    Returns a summary dict: {"imported": int, "skipped_duplicates": int, "total_in_file": int}
    """
    ext = os.path.splitext(path)[1].lower()

    if ext == ".csv":
        with open(path, "r", encoding="utf-8-sig", errors="ignore") as f:
            new_leads = parse_csv_text(f.read())
    elif ext == ".txt":
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            new_leads = parse_plain_lines(f.read())
    elif ext == ".docx":
        new_leads = parse_docx_tables(path)
    elif ext == ".pdf":
        new_leads = parse_pdf(path)
    else:
        raise ValueError(f"Unsupported file type: {ext}. Use .csv, .txt, .docx, or .pdf.")

    # De-dupe within this file's results first (e.g. same line matched twice)
    seen = set()
    deduped = []
    for lead in new_leads:
        key = lead.get("business_email") or lead.get("phone")
        if key and key not in seen:
            seen.add(key)
            deduped.append(lead)

    # Load existing leads.json and merge, skipping anything already present
    existing = []
    try:
        with open(LEADS_FILE, "r") as f:
            existing = json.load(f)
    except (FileNotFoundError, json.JSONDecodeError):
        existing = []

    existing_keys = {l.get("business_email") or l.get("phone") for l in existing if isinstance(l, dict)}

    imported_count = 0
    skipped_count = 0
    for lead in deduped:
        key = lead.get("business_email") or lead.get("phone")
        if key in existing_keys:
            skipped_count += 1
            continue

        lead.setdefault("niche", default_niche)
        lead.setdefault("preferred_channel", default_preferred_channel)
        lead["status"] = lead.get("status", "new")
        lead["source"] = "manual_import"
        lead["outreach_ready"] = bool(lead.get("business_email"))
        lead["needs_human"] = not lead.get("business_email")
        if lead["needs_human"]:
            lead["needs_human_reason"] = "Imported contact has no email — needs a call."

        existing.append(lead)
        existing_keys.add(key)
        imported_count += 1

    with open(LEADS_FILE, "w") as f:
        json.dump(existing, f, indent=2)

    return {
        "imported": imported_count,
        "skipped_duplicates": skipped_count,
        "total_in_file": len(deduped),
    }


if __name__ == "__main__":
    import sys
    if len(sys.argv) < 2:
        print("Usage: python contact_importer.py <path-to-file> [niche]")
        sys.exit(1)

    file_path = sys.argv[1]
    niche = sys.argv[2] if len(sys.argv) > 2 else ""
    result = import_contacts(file_path, default_niche=niche)
    print(f"Imported {result['imported']} new contacts "
          f"({result['skipped_duplicates']} duplicates skipped, "
          f"{result['total_in_file']} total found in file).")